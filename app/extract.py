from __future__ import annotations

import hashlib
import re
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any


@dataclass
class PageText:
    page: int  # 1-based
    text: str
    char_count: int = 0
    html: str = ""  # optional rich HTML for UI (tables preserved)
    tables: list[dict[str, Any]] = field(default_factory=list)
    # data:image/png;base64,... - bieu do/hinh nhung (khong phai full-page scan)
    images: list[str] = field(default_factory=list)
    ocr_applied: bool = False  # True neu trang duoc Gemini Vision lap text

    def __post_init__(self) -> None:
        self.char_count = len(self.text)


@dataclass
class DocumentText:
    path: str
    filename: str
    doc_id: str
    file_type: str
    pages: list[PageText] = field(default_factory=list)
    engine: str = ""
    total_pages: int = 0
    total_chars: int = 0
    sha256: str = ""
    warnings: list[str] = field(default_factory=list)

    def full_text(self, page_markers: bool = True) -> str:
        parts: list[str] = []
        for p in self.pages:
            if page_markers:
                parts.append(f"\n----- TRANG {p.page} -----\n{p.text}")
            else:
                parts.append(p.text)
        return "\n".join(parts).strip()

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


def _sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1 << 20), b""):
            h.update(chunk)
    return h.hexdigest()


def _clean_page_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _cell_clean(v: Any) -> str:
    if v is None:
        return ""
    s = str(v).replace("\x00", " ")
    s = re.sub(r"\s+", " ", s).strip()
    return s


def _matrix_to_markdown(rows: list[list[str]]) -> str:
    """Render a 2D cell matrix as GitHub-style markdown table."""
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    norm = [list(r) + [""] * (width - len(r)) for r in rows]
    norm = [r for r in norm if any(c.strip() for c in r)]
    if not norm:
        return ""
    keep_cols = [i for i in range(width) if any(r[i].strip() for r in norm)]
    if not keep_cols:
        return ""
    norm = [[r[i] for i in keep_cols] for r in norm]
    width = len(keep_cols)

    def esc(c: str) -> str:
        return c.replace("|", "\\|").replace("\n", " ")

    header = norm[0]
    body_start = 1
    if width and all(re.fullmatch(r"[\d.,%\-\s]+", c or "") for c in header if c):
        header = [f"Cot {i + 1}" for i in range(width)]
        body_start = 0
    lines = [
        "| " + " | ".join(esc(c) for c in header) + " |",
        "| " + " | ".join("---" for _ in header) + " |",
    ]
    for r in norm[body_start:]:
        lines.append("| " + " | ".join(esc(c) for c in r) + " |")
    return "\n".join(lines)


def _matrix_to_html(rows: list[list[str]], caption: str = "") -> str:
    """Compact, readable table HTML (inline styles - no frontend CSS edits)."""
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    norm = [list(r) + [""] * (width - len(r)) for r in rows]
    norm = [r for r in norm if any(c.strip() for c in r)]
    if not norm:
        return ""
    keep = [i for i in range(width) if any((r[i] or "").strip() for r in norm)]
    if not keep:
        return ""
    norm = [[r[i] for i in keep] for r in norm]

    wrap = (
        'style="max-width:100%;overflow-x:auto;margin:8px 0 12px;'
        'border:1px solid #333;background:#fff"'
    )
    table = (
        'style="border-collapse:collapse;width:100%;font-size:11.5px;'
        'line-height:1.35;font-family:\'Times New Roman\',serif"'
    )
    th = (
        'style="border:1px solid #333;padding:5px 8px;background:#fff;'
        'color:#000;font-weight:700;text-align:left;white-space:normal;'
        'font-size:11px"'
    )
    td = (
        'style="border:1px solid #333;padding:4px 8px;vertical-align:top;'
        'color:#000;word-break:break-word"'
    )
    cap = (
        'style="caption-side:top;text-align:left;padding:6px 8px 2px;'
        'font-size:11px;font-weight:600;color:#5b5f66"'
    )

    def cellhtml(c: str) -> str:
        # PyMuPDF to_markdown() dat "<br>" cho xuong dong trong o -> giu xuong dong
        return _html_esc(c).replace("&lt;br&gt;", "<br>").replace("&lt;br/&gt;", "<br>")

    parts = [f"<div {wrap}><table {table}>"]
    if caption:
        parts.append(f"<caption {cap}>{_html_esc(caption)}</caption>")
    parts.append("<thead><tr>")
    for c in norm[0]:
        parts.append(f"<th {th}>{cellhtml(c)}</th>")
    parts.append("</tr></thead><tbody>")
    for r in norm[1:]:
        parts.append("<tr>")
        for c in r:
            parts.append(f"<td {td}>{cellhtml(c)}</td>")
        parts.append("</tr>")
    parts.append("</tbody></table></div>")
    return "".join(parts)


def _html_esc(s: str) -> str:
    return (
        str(s)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def _bbox_overlap(
    a: tuple[float, float, float, float], b: tuple[float, float, float, float]
) -> bool:
    ax0, ay0, ax1, ay1 = a
    bx0, by0, bx1, by1 = b
    return not (ax1 < bx0 or bx1 < ax0 or ay1 < by0 or by1 < ay0)


def _page_has_ruling_lines(page: Any, *, min_h: int = 3, min_v: int = 1) -> bool:
    """Gate re: chi chay find_tables khi trang co duong ke bang.

    get_drawings() re hon find_tables rat nhieu -> bo qua trang toan chu.
    Bang thuc te (co vien) van duoc phat hien.
    """
    try:
        drawings = page.get_drawings()
    except Exception:
        return False
    h = v = 0
    for path in drawings:
        for it in path.get("items", []):
            kind = it[0]
            if kind == "l":  # line
                p1, p2 = it[1], it[2]
                if abs(p1.y - p2.y) < 1.0:
                    h += 1
                elif abs(p1.x - p2.x) < 1.0:
                    v += 1
            elif kind == "re":  # rectangle = 2 canh ngang + 2 canh doc
                h += 2
                v += 2
        if h >= min_h and v >= min_v:
            return True
    return h >= min_h and v >= min_v


def _extract_tables_pymupdf(page: Any) -> list[dict[str, Any]]:
    """PyMuPDF find_tables + to_markdown().

    Dung strategy="lines" de cat o theo DUONG VIEN THAT, khong theo canh chu
    (nguyen nhan "Din|h", "CO NG"). Bang cua tai lieu co vien nen rat hop.
    """
    out: list[dict[str, Any]] = []
    try:
        finder = page.find_tables(strategy="lines")
    except TypeError:
        try:
            finder = page.find_tables()
        except Exception:
            return out
    except Exception:
        return out
    tables = getattr(finder, "tables", None) or []
    for i, tab in enumerate(tables):
        matrix: list[list[str]] = []
        md = ""
        if hasattr(tab, "to_markdown"):
            try:
                md = str(tab.to_markdown() or "").strip()
            except Exception:
                md = ""
        try:
            raw = tab.extract() or []
            matrix = [[_cell_clean(c) for c in (row or [])] for row in raw]
        except Exception:
            matrix = []
        if not matrix and not md:
            continue
        if matrix and sum(1 for r in matrix for c in r if c) < 2:
            continue
        if not md and matrix:
            md = _matrix_to_markdown(matrix)
        if not md or "|" not in md:
            continue
        has_sep = any(
            len(row) > 1
            and all(re.fullmatch(r":?-{2,}:?", (c or "").strip() or "") for c in row)
            for row in [
                [c.strip() for c in ln.strip().strip("|").split("|")]
                for ln in md.splitlines()
                if ln.strip().startswith("|")
            ]
        )
        if not has_sep and matrix:
            md = _matrix_to_markdown(matrix)
        html = _matrix_to_html(matrix, caption=f"Bang {i + 1}") if matrix else ""
        if not html and md:
            rows_md = [
                [c.strip() for c in ln.strip().strip("|").split("|")]
                for ln in md.splitlines()
                if ln.strip().startswith("|")
                and not all(
                    re.fullmatch(r":?-{2,}:?", c.strip() or "")
                    for c in ln.strip().strip("|").split("|")
                )
            ]
            if rows_md:
                html = _matrix_to_html(rows_md, caption=f"Bang {i + 1}")
                matrix = rows_md
        bbox = tuple(
            float(x) for x in (tab.bbox if hasattr(tab, "bbox") else (0, 0, 0, 0))
        )
        out.append(
            {
                "index": i + 1,
                "bbox": bbox,
                "rows": len(matrix) if matrix else md.count("\n"),
                "cols": max((len(r) for r in matrix), default=0),
                "markdown": md,
                "html": html,
                "matrix": matrix,
                "source": "pymupdf_find_tables_lines",
            }
        )
    return out


def _reorder_masthead(
    blocks: list[tuple[float, float, float, float, str]],
    page_width: float,
) -> list[tuple[float, float, float, float, str]]:
    """Sap xep lai khoi tieu de (masthead) 2 cot cua van ban hanh chinh VN.

    PDF born-digital tra text theo thu tu doc tren->duoi nen masthead 2 cot bi
    DAN XEN: 'CO QUAN' / 'So: ...' (cot trai) tron voi 'CONG HOA...' /
    'Doc lap...' / 'ngay...' (cot phai). Ngoai ra Quoc hieu hay bi xuong dong
    ('...VIET' + 'NAM') nen regex Quoc hieu o frontend khong khop.

    Ham nay: tach 2 cot o vung dinh trang, gop dong Quoc hieu bi xuong dong,
    roi xuat theo thu tu chuan: cot trai (co quan, so hieu) -> cot phai
    (quoc hieu, tieu ngu, dia danh/ngay) -> phan than con lai giu nguyen.
    """
    if not blocks:
        return blocks
    TOP_ZONE = 165.0
    mid = (page_width or 595.0) / 2.0

    def cx(b: tuple[float, float, float, float, str]) -> float:
        return (b[0] + b[2]) / 2.0

    top = [b for b in blocks if b[1] < TOP_ZONE]
    rest = [b for b in blocks if b[1] >= TOP_ZONE]
    has_national = any(
        cx(b) >= mid and re.search(r"CỘNG\s*HÒA\s*XÃ\s*HỘI", b[4] or "")
        for b in top
    )
    if not has_national or len(top) < 2:
        return blocks

    left = sorted((b for b in top if cx(b) < mid), key=lambda b: b[1])
    right = sorted((b for b in top if cx(b) >= mid), key=lambda b: b[1])

    def flat(b: tuple[float, float, float, float, str]) -> str:
        return re.sub(r"\s+", " ", str(b[4]).replace("\n", " ")).strip()

    # Gop Quoc hieu bi xuong dong: '...VIET' + 'NAM' -> '...VIET NAM'
    merged_right: list[str] = []
    k = 0
    while k < len(right):
        t = flat(right[k])
        if re.search(r"CỘNG\s*HÒA", t) and not re.search(r"VIỆT\s*NAM", t):
            k += 1
            while k < len(right) and "NAM" not in t.upper():
                t = (t + " " + flat(right[k])).strip()
                k += 1
        else:
            k += 1
        if t:
            merged_right.append(t)

    ordered = [flat(b) for b in left if flat(b)] + merged_right
    y0 = min((b[1] for b in top), default=0.0)
    x0 = min((b[0] for b in top), default=0.0)
    synth = (x0, y0, page_width, TOP_ZONE, "\n".join(ordered))
    return [synth] + rest


def _page_text_excluding_tables(
    page: Any, table_bboxes: list[tuple[float, float, float, float]]
) -> str:
    """Extract body text while skipping table regions; reorder VN masthead.

    Dung get_text("blocks") de co toa do -> vua bo vung bang, vua sap xep lai
    khoi Quoc hieu / co quan ban hanh 2 cot (xem _reorder_masthead).
    """
    try:
        blocks = page.get_text("blocks") or []
    except Exception:
        return _clean_page_text(page.get_text("text") or "")

    kept: list[tuple[float, float, float, float, str]] = []
    for b in blocks:
        if len(b) < 5:
            continue
        x0, y0, x1, y1, text = b[0], b[1], b[2], b[3], b[4]
        if not str(text).strip():
            continue
        bb = (float(x0), float(y0), float(x1), float(y1))
        if any(_bbox_overlap(bb, tb) for tb in table_bboxes):
            continue
        kept.append((float(x0), float(y0), float(x1), float(y1), str(text)))
    if not kept:
        return _clean_page_text(page.get_text("text") or "")

    try:
        page_w = float(page.rect.width)
    except Exception:
        page_w = 595.0
    kept = _reorder_masthead(kept, page_w)
    parts = [str(b[4]).strip() for b in kept if str(b[4]).strip()]
    return _clean_page_text("\n".join(parts))


def _extract_page_images_from_page(
    doc: Any,
    page: Any,
    *,
    min_w: int = 90,
    min_h: int = 90,
    full_page_ratio: float = 0.82,
    max_px_width: int = 1000,
) -> list[str]:
    """Trich anh raster nhung tren 1 trang -> data-URL PNG.

    DUNG LAI doc/page dang mo (KHONG fitz.open() lai moi trang) -> het O(N)
    lan parse lai file. Bo icon/con dau nho va anh gan full-page (trang scan).
    """
    import base64

    import fitz

    page_area = float(page.rect.width * page.rect.height) or 1.0
    out: list[str] = []
    for img in page.get_images(full=True):
        xref = img[0]
        # Dien tich HIEN THI thuc te tren trang (point^2), KHONG phai pixel.
        # (So pixel khac don vi voi point -> ratio > 1 -> loai nham moi bieu do.)
        try:
            rects = page.get_image_rects(xref)
        except Exception:
            rects = []
        disp_area = 0.0
        for r in rects:
            disp_area = max(disp_area, abs(r.width * r.height))
        try:
            pix = fitz.Pixmap(doc, xref)
        except Exception:
            continue
        try:
            if pix.width < min_w or pix.height < min_h:
                continue
            # Chi loai anh CHIEM gan full-page (trang scan/nen) dua tren dien
            # tich HIEN THI, khong dua tren so pixel cua anh.
            if disp_area > 0 and disp_area > full_page_ratio * page_area:
                continue
            if pix.n - pix.alpha >= 4:  # CMYK -> RGB
                pix = fitz.Pixmap(fitz.csRGB, pix)
            png = pix.tobytes("png")
        finally:
            pix = None  # type: ignore[assignment]
        # Toi uu dung luong: ha be rong ve <= max_px_width + nen lai PNG.
        png = _downscale_png(png, max_px_width=max_px_width)
        out.append(
            "data:image/png;base64," + base64.b64encode(png).decode("ascii")
        )
    return out


def _downscale_png(png: bytes, *, max_px_width: int = 1000) -> bytes:
    """Ha be rong anh ve <= max_px_width va nen lai PNG (giam dung luong HTML)."""
    try:
        from PIL import Image
    except Exception:
        return png
    import io

    try:
        im = Image.open(io.BytesIO(png))
        if im.width > max_px_width:
            ratio = max_px_width / float(im.width)
            new_h = max(1, int(im.height * ratio))
            im = im.resize((max_px_width, new_h), Image.LANCZOS)
        buf = io.BytesIO()
        im.save(buf, format="PNG", optimize=True)
        data = buf.getvalue()
        return data if len(data) < len(png) else png
    except Exception:
        return png


def extract_page_images(
    pdf_path: str | Path,
    page_index: int,
    *,
    min_w: int = 90,
    min_h: int = 90,
    full_page_ratio: float = 0.82,
) -> list[str]:
    """Backward-compat wrapper (mo file 1 lan cho 1 trang).

    Duong dan chinh dung _extract_page_images_from_page voi doc dang mo.
    """
    import fitz

    doc = fitz.open(pdf_path)
    try:
        page = doc.load_page(page_index)
        return _extract_page_images_from_page(
            doc,
            page,
            min_w=min_w,
            min_h=min_h,
            full_page_ratio=full_page_ratio,
        )
    finally:
        doc.close()


def page_needs_ocr_local(text: str, *, min_chars: int = 40) -> bool:
    """Heuristic local (tranh import cycle voi vision_ocr luc load)."""
    t = (text or "").strip()
    if len(t) < min_chars:
        return True
    bad = sum(1 for c in t if ord(c) < 9 or c == "\ufffd")
    if bad > max(5, len(t) * 0.05):
        return True
    return False


def _extract_pdf_page(doc: Any, page: Any, page_no: int) -> PageText:
    """One PDF page: bang -> markdown + HTML; anh nhung dung lai doc dang mo.

    - Chi chay find_tables khi trang co duong ke (gate re) -> nhanh hon nhieu.
    - Da BO fallback aligned-lines (vua cham vua cat chu giua tu).
    """
    # Cong nhanh: chi chay find_tables khi trang co duong ke (rectangles/lines).
    # find_tables(strategy="lines") chi bat bang CO VIEN, nen gate nay KHONG bo
    # sot bang nao ma tranh clustering ton kem tren trang van xuoi -> nhanh hon.
    if _page_has_ruling_lines(page):
        tables = _extract_tables_pymupdf(page)
    else:
        tables = []

    bboxes = [tuple(t["bbox"]) for t in tables if t.get("bbox")]
    # Luon dung duong dan theo block (co toa do) de con sap xep lai masthead
    # 2 cot ngay ca khi trang KHONG co bang (trang dau moi van ban).
    body = _page_text_excluding_tables(page, bboxes)

    # Plain text cho LLM / search (giu markdown bang)
    text_parts: list[str] = []
    if body:
        text_parts.append(body)
    for t in tables:
        text_parts.append(
            f"\n[BANG {t.get('index', '')} - {t.get('rows')}x{t.get('cols')}]\n"
            f"{t.get('markdown', '')}\n"
        )

    # HTML cho UI
    html_parts: list[str] = []
    if body:
        for para in re.split(r"\n{2,}", body):
            p = para.strip()
            if not p:
                continue
            html_parts.append(f"<p>{_html_esc(p).replace(chr(10), '<br>')}</p>")
    for t in tables:
        html_parts.append(t.get("html") or "")

    pt = PageText(
        page=page_no,
        text="\n".join(text_parts).strip(),
        html="\n".join(html_parts),
        tables=[
            {
                "index": t.get("index"),
                "rows": t.get("rows"),
                "cols": t.get("cols"),
                "markdown": t.get("markdown"),
                "html": t.get("html"),
                "source": t.get("source", "find_tables"),
            }
            for t in tables
        ],
    )

    # Bieu do / anh nhung (chi trang born-digital) - DUNG LAI doc dang mo
    if not page_needs_ocr_local(pt.text):
        try:
            pt.images = _extract_page_images_from_page(doc, page)
        except Exception:
            pt.images = []
    return pt


def _extract_page_range_proc(
    path_str: str, idxs: list[int]
) -> list[tuple[int, PageText]]:
    """Worker chay trong TIEN TRINH rieng (an toan voi PyMuPDF).

    PyMuPDF/MuPDF khong thread-safe: nhieu THREAD (du moi thread mo doc rieng)
    van lam hong context toan cuc -> find_tables tra bbox/bang RAC. Moi TIEN
    TRINH co interpreter + context MuPDF rieng nen an toan tuyet doi.
    """
    import fitz

    wd = fitz.open(path_str)
    try:
        return [(i, _extract_pdf_page(wd, wd.load_page(i), i + 1)) for i in idxs]
    finally:
        wd.close()


def _extract_pdf(path: Path) -> tuple[list[PageText], str, list[str]]:
    warnings: list[str] = []
    try:
        import fitz  # PyMuPDF
    except ImportError as e:
        raise RuntimeError("PyMuPDF (pymupdf) is required for PDF extraction") from e

    import os
    import time

    base = fitz.open(path)
    pages: list[PageText] = []
    try:
        n = base.page_count
        _t0 = time.time()

        # PyMuPDF/MuPDF KHONG thread-safe: nhieu THREAD (du moi thread mo doc
        # rieng) van lam hong context toan cuc -> find_tables tra bbox/bang RAC
        # (bang "vo", tron van ban vao o header). Vi vay:
        #  - Mac dinh chay TUAN TU (an toan; ~3-7s cho 60-70 trang, thua <60s).
        #  - Dat EXTRACT_WORKERS>1 de bat song song hoa bang TIEN TRINH (process)
        #    -> moi tien trinh co interpreter rieng nen an toan.
        try:
            workers = int(os.environ.get("EXTRACT_WORKERS", "1") or 1)
        except ValueError:
            workers = 1
        if workers <= 0:
            workers = 1
        workers = max(1, min(workers, n or 1))

        pages = []
        used_parallel = False
        if workers > 1:
            try:
                import multiprocessing as _mp
                from concurrent.futures import ProcessPoolExecutor

                buckets: list[list[int]] = [[] for _ in range(workers)]
                for i in range(n):
                    buckets[i % workers].append(i)  # rai deu trang co bang
                nonempty = [b for b in buckets if b]
                collected: dict[int, PageText] = {}
                ctx = _mp.get_context("fork")
                with ProcessPoolExecutor(
                    max_workers=workers, mp_context=ctx
                ) as ex:
                    for part in ex.map(
                        _extract_page_range_proc,
                        [str(path)] * len(nonempty),
                        nonempty,
                    ):
                        for idx, pt in part:
                            collected[idx] = pt
                pages = [collected[i] for i in range(n)]
                used_parallel = True
            except Exception as e:  # noqa: BLE001 - fallback an toan
                warnings.append(
                    f"Song song hoa tien trinh loi ({e}); chay tuan tu."
                )
                pages = []

        if not used_parallel:
            pages = [
                _extract_pdf_page(base, base.load_page(i), i + 1) for i in range(n)
            ]

        workers_used = workers if used_parallel else 1
        warnings.append(
            f"Trich xuat {n} trang trong {round(time.time() - _t0, 1)}s "
            f"({workers_used} luong/tien trinh)."
        )

        empty = sum(1 for p in pages if not p.text)
        n_tables = sum(len(p.tables) for p in pages)
        if empty and empty == n:
            warnings.append(
                "PDF co ve scan/anh - text layer trong; se thu Gemini Vision OCR neu bat."
            )
        elif empty:
            warnings.append(f"{empty}/{n} trang khong co text (se OCR Vision neu bat).")
        if n_tables:
            warnings.append(f"Da giu cau truc {n_tables} bang (markdown/HTML).")

        # --- Gemini Vision OCR (PDF -> image -> text), parallel ---
        engine = "pymupdf+tables"
        try:
            import asyncio
            import concurrent.futures

            from .vision_ocr import (
                enrich_pages_with_vision,
                enrich_images_with_vision,
                ocr_enabled,
            )
            from .config import settings

            if ocr_enabled():
                mode = (settings.ocr_mode or "auto").lower()

                # Ngan sach thoi gian cho OCR (giay) -> bao ve muc tieu <60s.
                # Chinh bang OCR_TIMEOUT_S; dat <=0 de bo gioi han.
                try:
                    ocr_budget = float(os.environ.get("OCR_TIMEOUT_S", "30") or 30)
                except ValueError:
                    ocr_budget = 30.0

                async def _run_ocr():
                    async def _do():
                        pgs, w1 = await enrich_pages_with_vision(
                            path, pages, mode=mode
                        )
                        # Doc so lieu tu anh bieu do (chart) tren trang born-digital.
                        pgs, w2 = await enrich_images_with_vision(pgs, mode=None)
                        return pgs, (list(w1) + list(w2))

                    coro = _do()
                    if ocr_budget > 0:
                        return await asyncio.wait_for(coro, timeout=ocr_budget)
                    return await coro

                def _run_ocr_blocking():
                    # Luon dung event loop MOI -> an toan tren Python 3.13
                    # (khong dung asyncio.get_event_loop() da bi deprecated).
                    return asyncio.run(_run_ocr())

                try:
                    asyncio.get_running_loop()
                    # Dang trong 1 event loop -> day sang thread rieng de asyncio.run().
                    with concurrent.futures.ThreadPoolExecutor(max_workers=1) as pool:
                        pages, ocr_warn = pool.submit(_run_ocr_blocking).result()
                except RuntimeError:
                    # Khong co event loop dang chay -> chay truc tiep.
                    pages, ocr_warn = _run_ocr_blocking()

                warnings.extend(ocr_warn)
                if any("OCR Gemini" in w for w in ocr_warn):
                    engine = "pymupdf+tables+gemini-vision"
            elif empty:
                warnings.append(
                    "Bat OCR: set GEMINI_API_KEY + OCR_MODE=auto|always (gemini-2.0-flash Vision)."
                )
        except (asyncio.TimeoutError, TimeoutError):
            warnings.append(
                "OCR Vision qua han (OCR_TIMEOUT_S) - bo qua, giu nguyen text da trich de dam bao <60s."
            )
        except Exception as e:
            warnings.append(f"OCR Vision loi: {str(e)[:160]}")

        # --- Anh nhung: da trich inline trong _extract_pdf_page ---
        # Trang duoc OCR (scan) -> khong dan anh full-page.
        n_imgs = 0
        for pt in pages:
            if getattr(pt, "ocr_applied", False):
                pt.images = []
            else:
                n_imgs += len(pt.images or [])
        if n_imgs:
            warnings.append(f"Da trich {n_imgs} anh/bieu do nhung (raster).")

        return pages, engine, warnings
    finally:
        base.close()


def _extract_docx(path: Path) -> tuple[list[PageText], str, list[str]]:
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.text.paragraph import Paragraph
        from docx.table import Table
    except ImportError as e:
        raise RuntimeError("python-docx is required for Word extraction") from e

    document = Document(str(path))
    blocks: list[str] = []
    html_blocks: list[str] = []
    all_tables: list[dict[str, Any]] = []
    table_i = 0

    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            p = Paragraph(child, document)
            t = (p.text or "").strip()
            if t:
                blocks.append(t)
                html_blocks.append(f"<p>{_html_esc(t)}</p>")
        elif child.tag == qn("w:tbl"):
            table = Table(child, document)
            matrix: list[list[str]] = []
            for row in table.rows:
                cells: list[str] = []
                seen_ids: set[int] = set()
                for cell in row.cells:
                    cid = id(cell._tc)
                    if cid in seen_ids:
                        continue
                    seen_ids.add(cid)
                    cells.append(_cell_clean(cell.text))
                matrix.append(cells)
            if matrix and any(any(c for c in r) for r in matrix):
                table_i += 1
                md = _matrix_to_markdown(matrix)
                html = _matrix_to_html(matrix, caption=f"Bang {table_i}")
                blocks.append(
                    f"\n[BANG {table_i} - {len(matrix)}x{max(len(r) for r in matrix)}]\n{md}\n"
                )
                html_blocks.append(html)
                all_tables.append(
                    {
                        "index": table_i,
                        "rows": len(matrix),
                        "cols": max(len(r) for r in matrix),
                        "markdown": md,
                        "html": html,
                        "source": "docx",
                    }
                )

    full = "\n".join(blocks).strip()
    full_html = "\n".join(html_blocks)
    page_size = 2200
    pages: list[PageText] = []
    if not full:
        pages = [PageText(page=1, text="", html="", tables=[])]
    else:
        for i in range(0, len(full), page_size):
            chunk = full[i : i + page_size]
            t_in = [
                t
                for t in all_tables
                if t.get("markdown") and t["markdown"][:40] in chunk
            ]
            h_slice = full_html if i == 0 else ""
            pages.append(
                PageText(
                    page=len(pages) + 1,
                    text=chunk,
                    html=h_slice,
                    tables=t_in,
                )
            )
        if pages:
            assigned = {t["index"] for page in pages for t in page.tables}
            unassigned = [t for t in all_tables if t["index"] not in assigned]
            if unassigned:
                pages[0].tables = pages[0].tables + unassigned
            if not pages[0].html:
                pages[0].html = full_html

    return (
        pages,
        "python-docx+tables",
        [
            "Word: bang duoc giu dang markdown/HTML.",
            "So trang Word la uoc luong theo do dai.",
        ],
    )


def extract_document(path: str | Path) -> DocumentText:
    path = Path(path).resolve()
    if not path.exists():
        raise FileNotFoundError(str(path))

    suffix = path.suffix.lower()
    sha = _sha256_file(path)
    doc_id = sha[:16]

    if suffix == ".pdf":
        pages, engine, warnings = _extract_pdf(path)
        file_type = "pdf"
    elif suffix in (".docx", ".doc"):
        if suffix == ".doc":
            raise ValueError("Dinh dang .doc cu khong ho tro - chuyen sang .docx hoac PDF")
        pages, engine, warnings = _extract_docx(path)
        file_type = "docx"
    else:
        raise ValueError(f"Dinh dang khong ho tro: {suffix}. Dung PDF hoac DOCX.")

    total_chars = sum(p.char_count for p in pages)
    return DocumentText(
        path=str(path),
        filename=path.name,
        doc_id=doc_id,
        file_type=file_type,
        pages=pages,
        engine=engine,
        total_pages=len(pages),
        total_chars=total_chars,
        sha256=sha,
        warnings=warnings,
    )


def extract_folder(
    folder: str | Path,
    *,
    patterns: tuple[str, ...] = ("*.pdf", "*.docx", "*.PDF", "*.DOCX"),
    max_pages: int | None = None,
) -> list[DocumentText]:
    folder = Path(folder).resolve()
    if not folder.is_dir():
        raise NotADirectoryError(str(folder))

    files: list[Path] = []
    seen: set[Path] = set()
    for pat in patterns:
        for p in sorted(folder.glob(pat)):
            if p.is_file() and p not in seen:
                seen.add(p)
                files.append(p)

    docs: list[DocumentText] = []
    page_budget = max_pages
    for f in files:
        doc = extract_document(f)
        if page_budget is not None:
            if page_budget <= 0:
                break
            if doc.total_pages > page_budget:
                doc.pages = doc.pages[:page_budget]
                doc.total_pages = len(doc.pages)
                doc.total_chars = sum(p.char_count for p in doc.pages)
                doc.warnings.append(f"Cat con {page_budget} trang theo budget.")
                page_budget = 0
            else:
                page_budget -= doc.total_pages
        docs.append(doc)
    return docs


def chunk_pages(
    pages: list[PageText],
    *,
    pages_per_chunk: int = 8,  # legacy; ignored - chunk by char budget
    max_chars: int = 8000,
    source_file: str = "",
    max_chunks: int | None = None,
    hard_max_chunks: int | None = None,
) -> list[dict[str, Any]]:
    """Chia trang thanh chunk map-reduce theo nguong ky tu (khong cat giua trang)."""
    import math

    n = len(pages)
    if n == 0:
        return []

    cap = hard_max_chunks if hard_max_chunks is not None else max_chunks
    if cap is None or cap <= 0:
        cap = 40

    batches: list[list[PageText]] = []
    cur: list[PageText] = []
    cur_len = 0
    for p in pages:
        t = getattr(p, "text", "") or ""
        tlen = len(t)
        if cur and cur_len + tlen > max_chars:
            batches.append(cur)
            cur, cur_len = [], 0
        cur.append(p)
        cur_len += tlen
    if cur:
        batches.append(cur)

    if len(batches) > cap:
        step = math.ceil(len(batches) / cap)
        merged: list[list[PageText]] = []
        for i in range(0, len(batches), step):
            group: list[PageText] = []
            for b in batches[i : i + step]:
                group.extend(b)
            if group:
                merged.append(group)
        batches = merged

    chunks: list[dict[str, Any]] = []
    for batch in batches:
        body = "\n".join(f"[Trang {p.page}]\n{p.text or ''}" for p in batch)
        chunks.append(
            {
                "chunk_id": f"{source_file}:{batch[0].page}-{batch[-1].page}",
                "source_file": source_file,
                "page_start": batch[0].page,
                "page_end": batch[-1].page,
                "text": body,
                "char_count": len(body),
            }
        )
    return chunks
